from dataclasses import dataclass

import docx


@dataclass
class Subject:
    time: str
    title: list


@dataclass
class Day:
    title: str
    subjects: list


@dataclass
class subForGroup:
    time: str
    title: str


def fillDay(groups, dayTmp, row, indexes):
    for j in range(len(groups)):
        if len(dayTmp[j].subjects) != 0:
            if row.cells[1].text == dayTmp[j].subjects[-1].time:
                dayTmp[j].subjects[-1].title.append(row.cells[indexes[j]].text)
            else:
                dayTmp[j].subjects.append(Subject(row.cells[1].text, [row.cells[indexes[j]].text]))
        else:
            dayTmp[j].subjects.append(Subject(row.cells[1].text, [row.cells[indexes[j]].text]))


def updateGroups(groups, groups_first, groups_second, dayTmp):
    for j in range(len(groups)):
        groups_first[j][2].append(Day(dayTmp[j].title, []))
        groups_second[j][2].append(Day(dayTmp[j].title, []))
    for j in range(len(groups)):
        for k in range(len(dayTmp[j].subjects)):
            time = dayTmp[j].subjects[k].time
            sub = dayTmp[j].subjects[k].title
            groups_first[j][2][-1].subjects.append(subForGroup(time, sub[0]))
            groups_second[j][2][-1].subjects.append(subForGroup(time, sub[-1]))


def main():
    groups, indexes = [], []
    groupsRaw = docx.Document('schedules/АБ-98,99.docx').tables[1].rows[1]
    for i in range(2, len(groupsRaw.cells)):
        if groupsRaw.cells[i].text not in groups:
            groups.append(groupsRaw.cells[i].text)
            indexes.append(i)
    groups_first = [[groups[i], indexes[i], []] for i in range(len(groups))]
    groups_second = [[groups[i], indexes[i], []] for i in range(len(groups))]

    timesRaw = docx.Document('schedules/АБ-98,99.docx').tables[1]
    day = 'Понедельник'
    dayTmp = [Day(day, []) for i in range(len(groups))]
    for i in range(2, len(timesRaw.rows)):
        row = timesRaw.rows[i]
        if day == row.cells[0].text:
            fillDay(groups, dayTmp, row, indexes)
        else:
            updateGroups(groups, groups_first, groups_second, dayTmp)
            day = row.cells[0].text
            dayTmp = [Day(day, []) for i in range(len(groups))]
            fillDay(groups, dayTmp, row, indexes)
    updateGroups(groups, groups_first, groups_second, dayTmp)
    print()


if __name__ == '__main__':
    main()
